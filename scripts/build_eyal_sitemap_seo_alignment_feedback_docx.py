#!/usr/bin/env python3
"""Build Hebrew Word feedback for Eyal: sitemap + SEO/AEO/GEO alignment."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx_rtl import add_md_file_as_docx, add_rtl_paragraph, set_run_font


DOC_DATE = "2026-04-04"
VERSION = "v1"


def _paths():
    root = Path(__file__).resolve().parents[1]
    base = root / "docs/project/eyal-ceo-submissions-and-responses/to-eyal"
    wave = base / f"{DOC_DATE}--sitemap-seo-aeo-geo-feedback"
    md_source = wave / "md-sources" / f"{DOC_DATE}--eyal-sitemap-seo-feedback-response-he.md"
    out_dir = wave
    return root, md_source, out_dir


def main():
    _, md_source, out_dir = _paths()
    out_dir.mkdir(parents=True, exist_ok=True)
    out_docx = (
        out_dir
        / f"{DOC_DATE}--sitemap-seo-aeo-geo-alignment-feedback-for-eyal--{VERSION}.docx"
    )

    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    title_run = title.add_run("תשובת צוות — יישור עץ אתר ו־SEO · AEO · GEO")
    set_run_font(title_run, size=18, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sub_run = subtitle.add_run("אייל עמית · מסמך משוב להמשך תיאום")
    set_run_font(sub_run, size=12, bold=True)

    add_rtl_paragraph(
        doc,
        f"תאריך: {DOC_DATE} | גרסה: {VERSION} | מבוסס על מפת v2.3 מאושרת, אפיון סופי §4 ומצב M2",
        size=10,
    )
    doc.add_paragraph()

    add_md_file_as_docx(doc, md_source, skip_first_h1=True)

    doc.add_page_break()
    add_rtl_paragraph(doc, "הערות / החתמות", heading_level=1)
    add_rtl_paragraph(
        doc,
        "מקום להערות בכתב יד או בתוספת מסמך — ________________________________________________",
        size=11,
    )

    doc.save(out_docx)
    print(out_docx)


if __name__ == "__main__":
    main()
