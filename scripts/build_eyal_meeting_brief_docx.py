#!/usr/bin/env python3
"""מסמך Word לפגישה עם אייל: נעילות + מפת אתר + שאלות פתוחות עם אופציות והערות.

פלט: docs/project/eyal-ceo-submissions-and-responses/for-eyal/YYYY-MM-DD--meeting-brief/
Requires: pip install python-docx
"""
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx_rtl import add_md_file_as_docx, add_rtl_paragraph, set_run_font


def _paths():
    root = Path(__file__).resolve().parents[1]
    team100 = root / "docs/project/team-100-preplanning"
    eyal = root / "docs/project/eyal-ceo-submissions-and-responses"
    return root, team100, eyal


def _add_notes_lines(doc: Document, n: int = 3):
    for _ in range(n):
        add_rtl_paragraph(doc, "_________________________________________________________________")


def _add_question(
    doc: Document,
    num: int,
    title: str,
    context: str,
    options: tuple,
):
    add_rtl_paragraph(doc, f"שאלה {num} — {title}", heading_level=2)
    add_rtl_paragraph(doc, f"קונטקסט: {context}", size=10)
    add_rtl_paragraph(doc, "אפשרויות (סמן):", bold=True, size=10)
    for i, opt in enumerate(options, start=1):
        label = chr(ord("א") + i - 1) if i <= 3 else str(i)
        add_rtl_paragraph(doc, f"☐ ({label}) {opt}")
    add_rtl_paragraph(doc, "☐ שילוב / אחר — לפרט:", size=10)
    _add_notes_lines(doc, 2)


def main():
    _, team100, eyal = _paths()
    d = date.today().strftime("%Y-%m-%d")
    out_dir = eyal / "for-eyal" / f"{d}--meeting-brief"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_docx = out_dir / f"{d}--meeting-brief-for-eyal--v1.docx"

    doc = Document()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = t.add_run("פגישה — אייל עמית · תזכיר מסודר")
    set_run_font(r, size=20, bold=True)
    add_rtl_paragraph(doc, f"תאריך מסמך: {d}  |  מקור: אפיון סופי + נעילת v1.2 + טופס בחירות", size=10)
    doc.add_paragraph()

    # —— חלק 1 ——
    add_rtl_paragraph(doc, "1. קונספטים והחלטות שננעלו (סיכום לפגישה)", heading_level=1)
    locked = """• פלטפורמה: WordPress בלבד — כל התוכן בתוך WordPress; "רזה" = התקנה נקייה, תבנית רזה, מינימום תוספים מוקפים, ניקוי הפניות וקישורים שבורים.
• מקור אמת לתוכן: עריכה יומיומית ב־WordPress; Git — לא נדרש אישור CEO.
• נגישות: ללא יועץ נגישות חיצוני; עמידה בחוק ובתקן לאתר פרטי קטן.
• קוקיז / הסכמה בכניסה: דרישה מאושרת (מדיניות + מימוש טכני).
• QR: שימור כל URL מודפס; אינדקס QR — למנהל האתר בלבד (לא ציבורי).
• מסחר: ללא עגלה וסליקה ב־WP; כפתור לחשבונית ירוקה; קטלוג ראשי — שמירת slug קיים, תוכן מינימלי (לא 301 תחתון).
• IA: מותג אחד דומיין אחד; ספרים+הוצאה ממוזגים (בלי דף הוצאה נפרד); דפי ספר ל־SEO — מאושר; הופעות באזור משני "ארכיון"; קורסים — בתפריט ראשי לסקולר; הכשרות למטפלים — עמוד נפרד + "יעלה בקרוב"; הרצאות כמוצר — חלק מהמפה; בלוג מחודש.
• גלריה: נדרשת, פשוטה; נקודת פתיחה — העתקת מחיצות קיימות; הצוות מציע מיקום — לאישורך.
• 301 מפורט וסריקת קישורים שבורים — בשלבים מאוחרים יותר (לא חוסמים עכשיו).
• Google Analytics: צורך בקוד חדש (GA4) — בתוך פאזת עלייה / אופטימיזציה."""
    for line in locked.strip().split("\n"):
        add_rtl_paragraph(doc, line.strip())

    doc.add_page_break()

    # —— חלק 2 ——
    add_rtl_paragraph(doc, "2. מפת אתר עדכנית (טיוטה v2.1 — ממתין אישור)", heading_level=1)
    add_rtl_paragraph(
        doc,
        "להלן תוכן הקובץ SITEMAP-NEW-SITE-v2-DRAFT.md כפי שבמאגר. אחרי החלטות בסעיף 3 — נעדכן את המפה ל־APPROVED.",
        size=10,
    )
    sitemap_md = team100 / "SITEMAP-NEW-SITE-v2-DRAFT.md"
    add_md_file_as_docx(doc, sitemap_md, skip_first_h1=True)

    doc.add_page_break()

    # —— חלק 3 ——
    add_rtl_paragraph(doc, "3. שאלות פתוחות — סמן, בחר אפשרות, הערות", heading_level=1)
    add_rtl_paragraph(
        doc,
        "לכל שאלה: סמן באפשרות המתאימה או כתוב שילוב; השלם הערות בשורות.",
        size=10,
    )
    doc.add_paragraph()

    _add_question(
        doc,
        1,
        "ניסוח מותג ראשי (במקום «סטודיו» כשם גנרי)",
        "האתר משקף מרכז לטיפול ולימוד דיג'רידו. נדרש ניסוח ל־H1, תפריט ומטא־תיאורים.",
        (
            "המרכז לטיפול בדיג'רידו, נשימה מעגלית וסאונד הילינג — מלא.",
            "המרכז לטיפול בדיג'רידו — בשורה מתחת: נשימה מעגלית · סאונד הילינג · לימוד נגינה.",
            "אייל עמית — טיפול ולימוד דיג'רידו — שם אישי בחזית.",
        ),
    )
    _add_question(
        doc,
        2,
        "חשבונית ירוקה (Morning) — מסלול סליקה ראשי",
        "באתר: עמודי מוצר עם כפתור ליעד חיצוני. המחקר בצוות הושלם (Best+ לרוב קישורי תשלום).",
        (
            "קישורי תשלום מ־Morning (סכום קבוע / פתוח / הוראת קבע) מכל עמוד מוצר.",
            "דף מכירה / נחיתה hosted ב־Morning; באתר קישורים לשם.",
            "שילוב מפורש בין לינק ישיר לחלק מהמוצרים ודף מכירה לאחרים (לפרט בהערות).",
        ),
    )
    _add_question(
        doc,
        3,
        "תווית תפריט — מחיצה 1 (המרכז / השירותים הראשיים)",
        "פריט ראשי שמוביל לליבת העשייה (טיפול, לימוד, שיטה).",
        ("המרכז", "טיפול ולימוד", "מה אנחנו עושים"),
    )
    _add_question(
        doc,
        4,
        "תווית תפריט — מחיצה 2 (ספרים; הוצאה מוזגת)",
        "ללא דף הוצאה נפרד; מדור ספרים מרכזי.",
        ("ספרים", "ספרים ועוד ממני", "הוצאה וספרים"),
    )
    _add_question(
        doc,
        5,
        "תווית תפריט — מחיצה 3 (הופעות ועבר; לא חזית שיווקית)",
        "תוכן פעיל אבל בניווט משני; תווית שמשקפת 'עבר' בלי להסתיר.",
        ("הופעות וארכיון", "מהעבר שלי", 'ארכיון (מוסבר: "ישן יותר" אך באתר)'),
    )
    _add_question(
        doc,
        6,
        "עמוד אנגלית — כיוון",
        "פחות תיירות קלאסית; יותר קהילה בין־לאומית ומקצוע.",
        (
            "Eyal Amit — Didgeridoo, Breath & Sound",
            "Welcome — The Didgeridoo Center (Israel)",
            "Workshops & Sessions in English",
        ),
    )
    _add_question(
        doc,
        7,
        "דף הבית — כיוון עיצובי (תרשימים ב־HOME-PAGE-DIRECTIONS)",
        "מינימליסטי: תמונה גדולה, בלוק עדכני בולט, מעט אלמנטים משניים.",
        (
            "שער אחד — Hero + בלוק עדכני יחיד + עד 3 כרטיסים",
            "מרכז + עומק — Hero צר + עמודה + עדכונים + 2 בלוקים",
            "תמונה דומיננטית — כמעט מסך מלא + שורת חדשות + 2 פעולות",
        ),
    )
    add_rtl_paragraph(doc, "← תרשים ויזואלי לשאלה 7", heading_level=2)
    add_rtl_paragraph(
        doc,
        "קובץ HTML (דפדפן): for-eyal/assets/home-directions-visual.html — שלושה מסכים עם דוגמאות מדיה מהאתר הקיים (תיקיית home-preview). מומלץ לפתוח במקביל לפגישה או להצגה על מסך.",
        size=10,
    )
    add_rtl_paragraph(doc, "_" * 50, size=9)
    _add_question(
        doc,
        8,
        "קטלוג ראשי — שמירת נתיב (slug) קיים",
        "הוחלט בצוות: לא 301 תחתון במקום עמוד; תוכן מינימלי יותר. לאישור קריאה.",
        ("מאשר/ת שקראתי ומסכים/ה עם העיקרון", "נדרשת התאמנות — לפרט בהערות"),
    )
    _add_question(
        doc,
        9,
        "מפת אתר (סעיף 2 במסמך זה)",
        "האם העץ המוצג משקף את רצונך אחרי הבחירות למעלה? מה לשנות בכותרות או בסדר?",
        ("מאשר/ת את העקרונות והעץ כפי שמוצג", "שינויים נדרשים — לפרט בהערות"),
    )
    _add_question(
        doc,
        10,
        "גלריות (ארבע מחיצות קיימות באתר הנוכחי)",
        "הצוות יציע מיקום בעץ ובתפריט; נדרש אישורך אחרי ההצעה.",
        ("מאשר/ת את גישת 'העתקת מחיצות + הצעת מיקום'", "רוצה להגדיר מראש — לפרט בהערות"),
    )

    add_rtl_paragraph(doc, "הערות כלליות לפגישה", heading_level=2)
    _add_notes_lines(doc, 5)

    doc.save(out_docx)
    print(out_docx)


if __name__ == "__main__":
    main()
