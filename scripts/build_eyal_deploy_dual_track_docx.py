#!/usr/bin/env python3
"""Word (.docx) for Eyal: infrastructure policy — dual tracks, uPress, Git/deploy boundary.

Output: docs/project/eyal-ceo-submissions-and-responses/to-eyal/YYYY-MM-DD--infrastructure-policy-for-eyal/
Requires: pip install python-docx
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx_rtl import add_rtl_paragraph, set_run_font


DOC_DATE = "2026-03-29"
VERSION = "v1.0"


def _paths():
    root = Path(__file__).resolve().parents[1]
    eyal = root / "docs/project/eyal-ceo-submissions-and-responses"
    return root, eyal


def _add_table_2col(doc: Document, headers: tuple, rows: list[tuple[str, str]]):
    from docx.enum.table import WD_TABLE_ALIGNMENT

    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(h)
        set_run_font(r, bold=True)
    for ri, (a, b) in enumerate(rows, start=1):
        row = table.rows[ri].cells
        for ci, val in enumerate((a, b)):
            p = row[ci].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_run_font(p.add_run(val))


def main():
    _, eyal = _paths()
    out_dir = eyal / "to-eyal" / f"{DOC_DATE}--infrastructure-policy-for-eyal"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_docx = out_dir / f"{DOC_DATE}--infrastructure-dual-track-policy-for-eyal--{VERSION}.docx"

    doc = Document()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = t.add_run("תשתית, סביבות ומסלולי עבודה — לידיעת אייל עמית")
    set_run_font(r, size=18, bold=True)
    add_rtl_paragraph(doc, f"גרסה: {VERSION}  |  תאריך: {DOC_DATE}  |  מקור החלטות: צוות 100", size=10)
    add_rtl_paragraph(
        doc,
        "מסמך זה מסכם כיצד הצוות בונה ומפרסם את האתר (מקומי → סטייג'ינג → פרודקשן ב־uPress), "
        "ומפריד בין עבודה טכנית לבין תוכן ועיצוב שתלויים בך. המקור המלא לצוות נשאר במאגר (Markdown) — "
        "לאייל מועבר רק Word או PDF.",
        size=10,
    )
    doc.add_paragraph()

    add_rtl_paragraph(doc, "1. שני מסלולים במקביל", heading_level=1)
    add_rtl_paragraph(
        doc,
        "כל פיתוח האתר רץ במקביל על שני מסלולים. החוק המחייב: התשתית והקוד (מסלול ב׳) תמיד לפני מילוי תוכן סופי ועיצוב (מסלול א׳).",
    )
    _add_table_2col(
        doc,
        ("מסלול", "משמעות"),
        [
            (
                "א׳ — תוכן ועיצוב",
                "טקסטים, ויזואל, אישורים שלך, מילוי עשיר. תלוי בך; בזרם הזה הפרויקט נשאר ביחס לשלב M1 עד שתשחרר תוכן/כיוונים.",
            ),
            (
                "ב׳ — טכני, תשתית, קוד",
                "סביבת Docker מקומית, סטייג'ינג ב־uPress, תבנית, תוספים במינימום, Git, תיעוד תשתית. מתקדם לשלב M2 — ללא חסימה מתוכן.",
            ),
        ],
    )
    doc.add_paragraph()

    add_rtl_paragraph(doc, "2. סדר סביבות", heading_level=1)
    add_rtl_paragraph(doc, "• קודם: פיתוח מקומי בתוך Docker עד שיש בסיס יציב.")
    add_rtl_paragraph(doc, "• אחר כך: סטייג'ינג ייעודי ב־uPress (אחרי שהמקומי מוכן להעתקה או דחיפת קבצים).")
    add_rtl_paragraph(doc, "• פרודקשן: הדומיין והאתר הקיים כבר ב־uPress; מעבר סופי (cutover) לפי חבילת מיגרציה מאושרת.")
    doc.add_paragraph()

    add_rtl_paragraph(doc, "3. גרסת PHP ותאימות", heading_level=1)
    add_rtl_paragraph(
        doc,
        "נבחרת גרסת PHP אחת לפי מה ש־uPress מאפשר לחשבון. אותה גרסה משמשת גם בסביבה המקומית — כדי שלא יהיו הפתעות בין המחשב לשרת. צוות התשתית מתעד זאת ב־runbook.",
    )
    doc.add_paragraph()

    add_rtl_paragraph(doc, "4. Git לעומת מה שעולה לאתר", heading_level=1)
    add_rtl_paragraph(
        doc,
        "כל קוד האתר שנכנס ל־Git חי בעץ פריסה נפרד (תיקיית site — בעיקר wp-content: תבנית־בת, תוספים מותאמים אם יש). "
        "תיקיות תקשורת צוותים, מסמכי docs וסקריפטים — נשארים במאגר ולא מועלים לשרת.",
    )
    doc.add_paragraph()

    add_rtl_paragraph(doc, "5. דוא\"ל", heading_level=1)
    add_rtl_paragraph(doc, "לא מגדירים שליחת דוא\"ל בסביבה המקומית. הגדרות דוא\"ל בשרת — כשהסביבה באוויר, לפי uPress או תוסף.")
    doc.add_paragraph()

    add_rtl_paragraph(doc, "6. סיסמאות", heading_level=1)
    add_rtl_paragraph(
        doc,
        "סיסמאות וגישות מתועדות אצל צוות התשתית, מחוץ ל־Git. אחרי עלייה לאוויר — בוצעת רוטציה/עדכון סיסמאות למערכת הרלוונטית, כחלק מצ'קליסט המסירה.",
    )
    doc.add_paragraph()

    add_rtl_paragraph(doc, "7. קטלוג (כתובת חיה)", heading_level=1)
    add_rtl_paragraph(
        doc,
        "שומרים את נתיב הקטלוג כפי שהוא באתר החי היום (כולל QR ומודפסים) — בלי שינוי slug בלי החלטה מפורשת של ההנהלה.",
    )
    doc.add_paragraph()

    add_rtl_paragraph(doc, "8. גיבויים", heading_level=1)
    add_rtl_paragraph(
        doc,
        "ב־uPress משתמשים בגיבוי המובנה בפאנל. במחשב המקומי — Git לקוד. אין דרישה נפרדת ל־snapshot מקומי של מסד הנתונים לפרויקט זה.",
    )
    doc.add_paragraph()

    add_rtl_paragraph(doc, "9. מה נדרש עכשיו מהצוות הטכני (סיכום)", heading_level=1)
    for line in (
        "• לנעול גרסת PHP מול uPress וליישר את סביבת Docker.",
        "• להגדיר בבירור מה בדיוק עולה בפריסה לעומת מה נשאר רק במאגר.",
        "• לפתוח סטייג'ינג ב־uPress אחרי שהמקומי עובד.",
        "• לכלול במסירה: רוטציית סיסמאות אחרי go-live.",
    ):
        add_rtl_paragraph(doc, line)

    doc.save(out_docx)
    print(f"Wrote: {out_docx}")


if __name__ == "__main__":
    main()
