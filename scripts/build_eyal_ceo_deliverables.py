#!/usr/bin/env python3
"""Build Word (.docx) CEO packages for Eyal: executive summary, site map, decisions.

Outputs go to docs/project/eyal-ceo-submissions-and-responses/to-eyal/
and executive is also mirrored to team-100-preplanning/ for scripts compatibility.

Also writes dated folder: to-eyal/{DELIVERY_DATE}--final-spec-package-for-eyal/

Requires: pip install python-docx
"""
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx_rtl import add_md_file_as_docx, add_rtl_paragraph, set_run_font


DELIVERY_DATE = "2026-03-30"
VERSION_TAG = "v2-final"


def _paths():
    root = Path(__file__).resolve().parents[1]
    to_eyal = root / "docs/project/eyal-ceo-submissions-and-responses/to-eyal"
    team100 = root / "docs/project/team-100-preplanning"
    return root, to_eyal, team100


def build_executive_summary_docx():
    _, to_eyal, team100 = _paths()
    out_primary = to_eyal / f"{DELIVERY_DATE}--executive-summary--{VERSION_TAG}.docx"
    out_mirror = team100 / "EYAL-EXECUTIVE-SUMMARY-FOR-EYAL.docx"

    doc = Document()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = t.add_run("תקציר מנהלים + טופס אישור — אייל עמית")
    set_run_font(r, size=18, bold=True)

    add_rtl_paragraph(doc, f"גרסה: 2.0 (מסונכרן אפיון סופי)  |  תאריך: {DELIVERY_DATE}")
    add_rtl_paragraph(doc, "הוגש על ידי: נימרוד (ליווי מיגרציה)")
    add_rtl_paragraph(doc, "סטטוס: ממתין חתימת אייל עמית", bold=True)
    doc.add_paragraph()

    add_rtl_paragraph(doc, "תקציר מנהלים", heading_level=1)

    body = """מטרה: לצאת מלופי התיקון של האתר הישן ולבנות אתר WordPress אחד, פשוט לתחזוקה, עם דגש על מבקרים, SEO, ונגישות לפי דרישות החוק בישראל לאתר פרטי קטן. המכירות והסליקה יבוצעו מחוץ לאתר באמצעות חשבונית ירוקה; באתר יוצגו עמודי מוצר/שירות פשוטים עם כפתור להפניה לסליקה.

החלטות שאושרו בתכנון (עם נימרוד) — לאישורך הסופי:

1. פלטפורמה (נעול): WordPress בלבד — כל התוכן בתוך WordPress; אין חלק מהאתר מחוץ ל-WordPress. "רזה" = התקנה נקייה, תבנית מודרנית רזה ותקינה, מינימום תוספים (מוקפים ונחוצים), ניקוי פלונטרים של הפניות וקישורים שבורים, ועדכון עץ האתר לפעילותך.

2. מותגים: מגוון מותגים באותו דומיין — הפרדה בחוויית משתמש בלבד (לא מערכות נפרדות).

3. מרכז: הסטודיו, הנשימה והדיג'רידו — המרכז בדף הבית ובתפריט הראשי.

4. הוצאה ומופע: חלק מהעבר, נשארים וחשובים — מוצגים כארכיון מותגי משני שמעשיר את האתר.

5. חנות: ללא עגלה ותשלום באתר; עמוד קטלוג ראשי עם שמירת נתיב קיים (לא 301 תחתון במקום עמוד); חשבונית ירוקה לסליקה.

6. עמודי QR: חובה לשמור את כל הכתובות (מודפסים בספרים). אסור לשנות slug או לבצע 301 ששובר סריקה. אינדקס QR — למנהל האתר בלבד (לא ציבורי).

7. בלוג: להחזיר לחיים כנכס SEO מרכזי.

8. אירועים: בלוק "אירוע הבא" בתבנית; טפסים — קצר באתר + מפורט חיצונית.

9. SEO מקומי: סכמות עסק מקומי (לפי מה שמאושר מקצועית).

10. נגישות: עמידה בחוק ובתקן לאתר פרטי קטן — ללא ייעוץ חיצוני (אישור צוות טכני).

11. תפעול: אייל — מנהל ומתחזק; נימרוד — ליווי עד יציבות.

12. שפה: עמוד נחיתה באנגלית בנוסף לעברית (כיוון בינלאומי; פרטים בטופס הבחירות).

מדדי הצלחה: לאזן יציבות, תחזוקה, SEO ו-UX — וגם תנועה אמיתית; בלי מבקרים היעד העסקי לא הושג."""

    for para in body.split("\n\n"):
        add_rtl_paragraph(doc, para.strip())

    doc.add_paragraph()
    add_rtl_paragraph(doc, "אישור עקרונות (סמן וחתום)", heading_level=1)

    table = doc.add_table(rows=8, cols=3)
    table.style = "Table Grid"
    hdr = ("#", "נושא", "מאושר (סמן V)")
    for i, h in enumerate(hdr):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(h)
        set_run_font(run, bold=True)
    rows_data = [
        ("1", "אתר WordPress אחד, ענף נקי, ארכיון מסלול קודם", ""),
        ("2", "הפרדת UX בין סטודיו / הוצאה+ספרים / מופעים", ""),
        ("3", "סליקה בחשבונית ירוקה בלבד (ללא עגלה באתר)", ""),
        ("4", "שימור מלא URLי QR ללא שבירת ספרים מודפסים", ""),
        ("5", "בלוג כנכס SEO מחודש", ""),
        ("6", "נגישות — עמידה בחוק ובתקן (ללא ייעוץ חיצוני — כפי שסוכם)", ""),
        ("7", "עמוד נחיתה באנגלית", ""),
    ]
    for r_idx, (num, topic, _) in enumerate(rows_data, start=1):
        row = table.rows[r_idx].cells
        for c_idx, val in enumerate((num, topic, "_______")):
            p = row[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(val)
            set_run_font(run)

    doc.add_paragraph()
    add_rtl_paragraph(doc, "שם: _________________________    חתימה: __________________    תאריך: __________")
    doc.add_paragraph()
    add_rtl_paragraph(doc, "הערות אייל:", bold=True)
    add_rtl_paragraph(doc, "_________________________________________________________________")
    add_rtl_paragraph(doc, "_________________________________________________________________")

    doc.add_paragraph()
    add_rtl_paragraph(doc, "פתוח למילוי על ידי אייל (לפני בנייה מלאה)", heading_level=1)

    open_items = """1. חשבונית ירוקה: אחרי מחקר יכולות — יוצגו אופציות ישימות; עד אז ראו GREEN-INVOICE-LINK-MAP במאגר.

2. אינדקס QR: נסגר טכנית — למנהל האתר בלבד.

3–5. מותג, תפריט (מחיצות), EN — טופס בחירות נפרד (קובץ for-eyal-choices) עם 3 אופציות לכל סעיף.

6. רשימת תוכן (Keep/Merge/Drop): הצוות מכין דוח; אם נדרש — ממשק HTML לסינון; אתה מאשר — לא מילוי שורה־שורה.

7. גלריות: מוצע מיקום בעץ ובתפריט בהגשה — לאישורך.

8. דיון 301 מפורט — בשלב מאוחר; קטלוג ראשי נשמר בנתיב קיים.

לאחר מילוי: להחזיר לנימרוד בערוץ העבודה."""

    for para in open_items.split("\n\n"):
        add_rtl_paragraph(doc, para.strip())

    doc.add_paragraph()
    add_rtl_paragraph(
        doc,
        "אפיון סופי למימוש (צוות): SITE-SPECIFICATION-FINAL-2026-03-30.md. מסמכים ישנים: LEGACY-DOCUMENTS-INDEX-2026-03-30.md.",
        size=9,
    )
    add_rtl_paragraph(
        doc,
        "חבילת קבצים להגשה: to-eyal/" + DELIVERY_DATE + "--final-spec-package-for-eyal/",
        size=9,
    )

    to_eyal.mkdir(parents=True, exist_ok=True)
    doc.save(out_primary)
    doc.save(out_mirror)
    return out_primary, out_mirror


def build_site_map_docx():
    _, to_eyal, team100 = _paths()
    out = to_eyal / f"{DELIVERY_DATE}--site-map-draft-v2--{VERSION_TAG}.docx"
    md_path = team100 / "SITEMAP-NEW-SITE-v2-DRAFT.md"

    doc = Document()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = t.add_run("מפת אתר — טיוטה לאישור אייל (v2)")
    set_run_font(r, size=18, bold=True)
    add_rtl_paragraph(doc, f"תאריך הגשה: {DELIVERY_DATE}  |  סטטוס: טיוטה — ממתין אישור")
    add_rtl_paragraph(
        doc,
        "מסמך זה משקף את מבנה המידע המוצע לאתר החדש. אחרי אישורך נעבור לאפיון עמודים מפורט.",
        size=10,
    )
    doc.add_paragraph()
    add_md_file_as_docx(doc, md_path, skip_first_h1=True)
    doc.add_paragraph()
    add_rtl_paragraph(doc, "חתימה לאישור מפת אתר: __________________  תאריך: __________", bold=True)
    to_eyal.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


def _add_table(doc, headers: tuple, rows: tuple):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        p = table.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(h)
        set_run_font(run, bold=True)
    for ri, row_vals in enumerate(rows, start=1):
        for ci, val in enumerate(row_vals):
            p = table.rows[ri].cells[ci].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(val)
            set_run_font(run)
    return table


def build_decisions_docx():
    _, to_eyal, team100 = _paths()
    out = to_eyal / f"{DELIVERY_DATE}--decisions-for-approval--{VERSION_TAG}.docx"

    doc = Document()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = t.add_run("קובץ החלטות לאישור אייל — אתר חדש eyalamit.co.il")
    set_run_font(r, size=18, bold=True)
    add_rtl_paragraph(doc, f"גרסה: 2.1 (מסונכרן אפיון סופי)  |  תאריך: {DELIVERY_DATE}")
    add_rtl_paragraph(
        doc,
        "מקור מחייב לצוות: SITE-SPECIFICATION-FINAL-2026-03-30.md. מסמכי 01–06 וטיוטות קודמות — LEGACY בלבד.",
        size=10,
    )
    doc.add_paragraph()

    add_rtl_paragraph(doc, "1. מהות האתר (שער 0)", heading_level=1)
    _add_table(
        doc,
        ("הצהרה", "נדרש אישור (סמן)"),
        (
            ("האתר הוא אתר שיווקי־מותגי: מידע, אמון, והנעה לפעולה", "[ ]"),
            ("האתר אינו מערכת מכירות מקוונת או עגלת קניות", "[ ]"),
            ("מורכבות (מכירה, תשלומים) — במערכות חיצוניות לפי צורך", "[ ]"),
        ),
    )
    doc.add_paragraph()
    add_rtl_paragraph(doc, "חתימה: __________________  תאריך: ________")

    doc.add_paragraph()
    add_rtl_paragraph(doc, "2. מדדי הצלחה (תמצית)", heading_level=1)
    _add_table(
        doc,
        ("מדד", "יעד מוצע"),
        (
            ("ביצועים מובייל", "Lighthouse Performance ≥ 85 (או שקיל)"),
            ("יציבות", "0 שגיאות קונסול קריטיות בעמודי ליבה"),
            ("תחזוקה", "רשימת תוספים מאושרת ומצומצמת"),
            ("SEO טכני", "sitemap אחד, ללא כפילויות קנוניקל חמורות"),
            ("בהירות מסר", "מבקר מבין מה נעשה, למי, איך מתחילים"),
        ),
    )
    add_rtl_paragraph(doc, "תקרת תקציב שנתית (למילוי): __________________")

    doc.add_paragraph()
    add_rtl_paragraph(doc, "3. פלטפורמה (נעול — אין בחירה חלופית)", heading_level=1)
    add_rtl_paragraph(
        doc,
        "WordPress בלבד: כל התוכן בפנים. רזה = התקנה נקייה, תבנית מודרנית רזה, מינימום תוספים מוקפים, ניקוי הפניות/קישורים שבורים, עדכון עץ האתר. אין יציאה של תוכן מחוץ ל-WordPress.",
    )
    add_rtl_paragraph(doc, "אישור התיישבות: ________  חתימה: __________________  תאריך: __________")

    doc.add_paragraph()
    add_rtl_paragraph(doc, "4. החלטות מתועדות (לסיכום — לאישורך)", heading_level=1)
    bullets = (
        "מותגים: אתר אחד, דומיין אחד; הפרדת UX בין סטודיו (מרכז), הוצאה/ספרים, מופעים (ארכיון).",
        "מקור אמת לתוכן: WordPress (אתר חי); Git — לא SSOT לאייל.",
        "מסחר: ללא עגלה ב-WordPress; סליקה בחשבונית ירוקה; קטלוג ראשי — שמירת נתיב קיים.",
        "QR: שימור כל URL מודפס; אינדקס QR — למנהל בלבד.",
        "בלוג: החייאה כנכס SEO.",
        "נגישות: חוק ותקן לאתר פרטי קטן — ללא יועץ חיצוני (אישור צוות).",
        "תפעול: אייל — מנהל; נימרוד — מיגרציה עד יציבות.",
        "אנגלית: עמוד נחיתה — פרטים בטופס בחירות.",
        "תוכן (KMD): דוח צוות + אישור אייל; אופציונלי ממשק HTML לסינון.",
        "גלריות: העתקת מחיצות קיימות + הצעת מיקום בתפריט לאישור.",
    )
    for b in bullets:
        add_rtl_paragraph(doc, "• " + b)

    doc.add_paragraph()
    add_rtl_paragraph(doc, "5. מדיניות תוכן (Keep / Merge / Drop) — עקרונות מחייבים", heading_level=1)
    add_rtl_paragraph(
        doc,
        "• QR: Keep לנתיב — אין שינוי slug ואין 301 (אלא אם תחליט אחרת בכתב).",
    )
    add_rtl_paragraph(
        doc,
        "• חנות Woo: Drop ממבנה האתר החדש + 301 ליעד שתאשר (לאחר בדיקת קישורים).",
    )
    add_rtl_paragraph(
        doc,
        "• בלוג: לפי כוונתך (החייאה ואופטימיזציה); סינון תוכן חלש — המלצה משנית.",
    )
    add_rtl_paragraph(
        doc,
        "רשימת תוכן מלאה: דוח שמכין הצוות; אייל מאשר — לא מילוי טכני שורה-שורה. (ממשק HTML אופציונלי לסינון.)",
        size=10,
    )

    doc.add_paragraph()
    add_rtl_paragraph(doc, "הערות אייל:", bold=True)
    add_rtl_paragraph(doc, "_________________________________________________________________")
    doc.add_paragraph()
    add_rtl_paragraph(doc, "מסמך מקור מחייב לצוות: SITE-SPECIFICATION-FINAL-2026-03-30.md", size=9)

    to_eyal.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


def build_for_eyal_choices_docx():
    _, to_eyal, _ = _paths()
    eyal_dir = to_eyal.parent
    md_path = eyal_dir / "FOR-EYAL-CHOICES-v1.2-2026-03-30.md"
    out = to_eyal / f"{DELIVERY_DATE}--for-eyal-choices-v1.2--{VERSION_TAG}.docx"

    doc = Document()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = t.add_run("טופס בחירות לאייל — אפיון אתר v1.2")
    set_run_font(r, size=18, bold=True)
    add_rtl_paragraph(doc, f"תאריך הגשה: {DELIVERY_DATE}  |  ייצוא אוטומטי מ-Markdown פנימי")
    add_rtl_paragraph(doc, "לסמן א / ב / ג או שילוב — ראו מסמך המקור בצוות אם נדרש עדכון.", size=10)
    doc.add_paragraph()
    add_md_file_as_docx(doc, md_path, skip_first_h1=True)
    doc.add_paragraph()
    add_rtl_paragraph(doc, "חתימה: __________________  תאריך: __________", bold=True)
    to_eyal.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


def build_green_invoice_action_sheet_docx():
    _, to_eyal, _ = _paths()
    eyal_dir = to_eyal.parent
    md_path = eyal_dir / "FOR-EYAL-GREEN-INVOICE-ACTION-SHEET-2026-03-30.md"
    out = to_eyal / f"{DELIVERY_DATE}--for-eyal-green-invoice-action-sheet--{VERSION_TAG}.docx"

    doc = Document()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = t.add_run("חשבונית ירוקה — תקציר, המלצה ובדיקות לאישור אייל")
    set_run_font(r, size=18, bold=True)
    add_rtl_paragraph(doc, f"תאריך הגשה: {DELIVERY_DATE}  |  ייצוא אוטומטי מ-Markdown פנימי")
    add_rtl_paragraph(
        doc,
        "לאחר מילוי הבדיקות והדוגמה — להחזיר לנימרוד. מחקר מלא אצל הצוות בלבד.",
        size=10,
    )
    doc.add_paragraph()
    add_md_file_as_docx(doc, md_path, skip_first_h1=True)
    doc.add_paragraph()
    add_rtl_paragraph(doc, "חתימה: __________________  תאריך: __________", bold=True)
    to_eyal.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


def build_site_spec_final_docx():
    _, to_eyal, team100 = _paths()
    md_path = team100 / "SITE-SPECIFICATION-FINAL-2026-03-30.md"
    out = to_eyal / f"{DELIVERY_DATE}--site-specification-final--{VERSION_TAG}.docx"

    doc = Document()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = t.add_run("אפיון אתר — גרסה סופית (מקור מחייב לצוות)")
    set_run_font(r, size=18, bold=True)
    add_rtl_paragraph(doc, f"תאריך: {DELIVERY_DATE}  |  עותק להגשה / לצוות; מסמכים ישנים — LEGACY-INDEX")
    doc.add_paragraph()
    add_md_file_as_docx(doc, md_path, skip_first_h1=True)
    doc.add_paragraph()
    add_rtl_paragraph(doc, "אישור צוות / הערות אייל (אם רלוונטי): __________________", bold=True)
    to_eyal.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


def assemble_final_package(files: tuple[Path, ...]) -> Path:
    _, to_eyal, _ = _paths()
    pkg = to_eyal / f"{DELIVERY_DATE}--final-spec-package-for-eyal"
    if pkg.exists():
        shutil.rmtree(pkg)
    pkg.mkdir(parents=True, exist_ok=True)
    for p in files:
        if p.is_file():
            shutil.copy2(p, pkg / p.name)
    for junk in ("Icon\r", "Icon"):
        j = pkg / junk
        if j.exists():
            j.unlink()
    readme = f"""חבילת אפיון סופית — {DELIVERY_DATE}
================================

לאייל עמית — Word (.docx) בלבד להגשה רשמית.
אל תשלחו קבצי .md לאייל (ראו SSOT).

קבצים בתיקייה זו:
"""
    for p in files:
        if p.is_file():
            readme += f"  - {p.name}\n"
    readme += """
מקורות Markdown פנימיים (צוות):
  team-100-preplanning/SITE-SPECIFICATION-FINAL-2026-03-30.md
  team-100-preplanning/GREEN-INVOICE-CAPABILITIES-FINDINGS-2026-03-30.md
  eyal-ceo-submissions-and-responses/FOR-EYAL-GREEN-INVOICE-ACTION-SHEET-2026-03-30.md

לייצוא PDF: פתחו כל .docx ב-Word ושמרו כ-PDF.
"""
    (pkg / "README.txt").write_text(readme, encoding="utf-8")
    return pkg


def main():
    e1, e2 = build_executive_summary_docx()
    s = build_site_map_docx()
    d = build_decisions_docx()
    c = build_for_eyal_choices_docx()
    g = build_green_invoice_action_sheet_docx()
    f = build_site_spec_final_docx()
    primary_outputs = (e1, s, d, c, g, f)
    pkg = assemble_final_package(primary_outputs)
    print("Written:")
    for p in primary_outputs:
        print(" ", p)
    print(" ", e2, "(mirror)")
    print("Package:")
    print(" ", pkg)


if __name__ == "__main__":
    main()
