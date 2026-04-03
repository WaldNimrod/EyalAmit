#!/usr/bin/env python3
"""Build a client-facing Hebrew Word report for Eyal: SEO + AEO + GEO readiness."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx_rtl import add_hyperlink, add_md_file_as_docx, add_rtl_paragraph, set_run_font


DOC_DATE = "2026-03-31"
VERSION = "v1"


def _paths():
    root = Path(__file__).resolve().parents[1]
    base = (
        root
        / "docs/project/eyal-ceo-submissions-and-responses/to-eyal"
        / f"{DOC_DATE}--aeo-geo-client-report"
    )
    md_source = base / "md-sources" / f"{DOC_DATE}-AEO-GEO-client-report-he.md"
    out_dir = base
    return root, md_source, out_dir


def _add_reference(doc: Document, title: str, url: str, note: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    lead = p.add_run(f"{title}: ")
    set_run_font(lead, bold=True)
    add_hyperlink(p, url, url)
    if note:
        tail = p.add_run(f" | {note}")
        set_run_font(tail, size=10)


def main():
    _, md_source, out_dir = _paths()
    out_dir.mkdir(parents=True, exist_ok=True)
    out_docx = out_dir / f"{DOC_DATE}--aeo-geo-client-report-for-eyal--{VERSION}.docx"

    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    title_run = title.add_run("דוח מסכם לאייל עמית")
    set_run_font(title_run, size=20, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    subtitle_run = subtitle.add_run("מוכנות האתר החדש ל־SEO, AEO ו־GEO")
    set_run_font(subtitle_run, size=16, bold=True)

    add_rtl_paragraph(
        doc,
        f"תאריך מסמך: {DOC_DATE} | מבוסס על בדיקת האתר החי, מסמכי הפרויקט והנחיות רשמיות עדכניות",
        size=10,
    )
    add_rtl_paragraph(
        doc,
        "מטרת המסמך: לספק בסיס מסודר לשיחה עם אייל, קבלת החלטות והמשך מחקר לפני build מלא.",
        size=10,
    )
    doc.add_paragraph()

    add_md_file_as_docx(doc, md_source, skip_first_h1=True)

    doc.add_page_break()

    add_rtl_paragraph(doc, "רפרנסים עיקריים", heading_level=1)
    _add_reference(
        doc,
        "Google Search Central: AI features and your website",
        "https://developers.google.com/search/docs/appearance/ai-features",
        "הבהרה רשמית של Google: אין דרישות טכניות מיוחדות מעבר ליסודות חיפוש תקינים.",
    )
    _add_reference(
        doc,
        "Google Search Central Blog: succeeding in AI search",
        "https://developers.google.com/search/blog/2025/05/succeeding-in-ai-search",
        "מיקוד בתוכן מועיל, crawl תקין וחוויית משתמש טובה גם בעולמות AI.",
    )
    _add_reference(
        doc,
        "OpenAI Crawlers Documentation",
        "https://developers.openai.com/api/docs/bots",
        "מגדיר את תפקידי OAI-SearchBot ו-GPTBot ואת ההפרדה ביניהם.",
    )
    _add_reference(
        doc,
        "OpenAI Publishers and Developers FAQ",
        "https://help.openai.com/en/articles/12627856-publishers-and-developers-faq",
        "כולל התייחסות למדידת תנועה מ-ChatGPT Search דרך utm_source=chatgpt.com.",
    )
    _add_reference(
        doc,
        "האתר החי של אייל עמית",
        "https://www.eyalamit.co.il/",
        "נבדק בפועל במסגרת האבחון.",
    )
    _add_reference(
        doc,
        "robots.txt באתר החי",
        "https://www.eyalamit.co.il/robots.txt",
        "נכון ל-2026-03-31 מחזיר 404.",
    )
    _add_reference(
        doc,
        "sitemap_index.xml באתר החי",
        "https://www.eyalamit.co.il/sitemap_index.xml",
        "זמין ומנוהל דרך Yoast.",
    )

    doc.add_paragraph()
    add_rtl_paragraph(doc, "המלצות למחקר נוסף", heading_level=1)
    add_rtl_paragraph(
        doc,
        "1. מחקר שאילתות בעברית ובאנגלית לפי כוונת חיפוש: שאלות טיפול, לימוד, נשימה, סדנאות וסאונד.",
    )
    add_rtl_paragraph(
        doc,
        "2. מחקר ישויות ומתחרים: איך גופים דומים מוצגים במנועי חיפוש, במפות, ובתוצאות AI.",
    )
    add_rtl_paragraph(
        doc,
        "3. מחקר מדידה: הגדרת KPI, מקורות תנועה, המרות, ודשבורד לתנועה אורגנית ותנועה מסביבות AI.",
    )
    add_rtl_paragraph(
        doc,
        "4. מחקר editorial: תכנית תוכן רבעונית לשילוב עמודי שירות, FAQ, מאמרי עומק ועדכוני אירועים.",
    )

    doc.add_paragraph()
    add_rtl_paragraph(doc, "הערות / החלטות אייל", heading_level=1)
    for _ in range(6):
        add_rtl_paragraph(doc, "_________________________________________________________________")

    doc.save(out_docx)
    print(out_docx)


if __name__ == "__main__":
    main()
