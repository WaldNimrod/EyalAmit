#!/usr/bin/env python3
"""Build meeting checklist .docx with clickable links (upload to Google Drive → Open as Google Doc)."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

HUB = "http://eyalamit-co-il-2026.s887.upress.link/ea-eyal-hub/"
STAGING = "http://eyalamit-co-il-2026.s887.upress.link"


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def rtl_paragraph(doc: Document, style: str | None = None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.right_to_left = True
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return p


def add_heading_rtl(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.right_to_left = True
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_bullet_rtl(doc: Document, text: str, checked: bool = False) -> None:
    prefix = "☐ " if checked is False else "☑ "
    p = rtl_paragraph(doc)
    p.add_run(prefix + text)


def add_link_line(doc: Document, label: str, url: str) -> None:
    p = rtl_paragraph(doc)
    p.add_run(label + " ")
    add_hyperlink(p, url, url)


def hub(path: str) -> str:
    return HUB + path.lstrip("/")


def staging(path: str) -> str:
    return STAGING + path


def build(out_path: Path) -> None:
    doc = Document()
    for section in doc.sections:
        section.page_height = section.page_height
        section.page_width = section.page_width

    add_heading_rtl(doc, "צ'קליסט פגישה — אייל · Chapters full-site · 2026-06-23", 0)

    p = rtl_paragraph(doc)
    run = p.add_run("לנמרוד בלבד (מסמך פנימי — לא להגשה לאייל כ-.md)")
    run.italic = True
    run.font.size = Pt(11)

    add_heading_rtl(doc, "קישורי בסיס", 2)
    add_link_line(doc, "מה נדרש ממך (עדיפית):", hub("what-we-need.html"))
    add_link_line(doc, "Staging (אתר הבדיקה):", STAGING + "/")
    p = rtl_paragraph(doc)
    p.add_run("גרסת Hub: 1.9.1 · בנייה 2026-06-23")

    add_heading_rtl(doc, "מנופים עסקיים (למה הפגישה חשובה)", 1)
    for line in [
        "CP-01 — פילר נחירות/דום נשימה: המנוף הגדול ביותר לתעבורה אורגנית ו-AI citation.",
        "WP-01 — מעקב המרות (generate_lead + wa.me) — כבר ב-Wave-1; אימות ב-cutover.",
        "48 עדויות — E-E-A-T ואמון בעמודי שירות.",
        "BN-01 — מותג/ישות: הסרת «סטודיו נשימה מעגלית» לעקביות NAP.",
    ]:
        add_bullet_rtl(doc, line)
    p = rtl_paragraph(doc)
    p.add_run("פרטים: ")
    add_hyperlink(p, hub("content-proposals.html"), hub("content-proposals.html"))

    add_heading_rtl(doc, "לפני הפגישה (5 דק)", 1)
    checks = [
        ("Hub — what-we-need.html: 7 עדיפויות", hub("what-we-need.html")),
        ("תדריך פגישה", hub("meeting.html")),
        ("Staging מוכן לסיור", staging("/?nc=1")),
    ]
    for label, url in checks:
        p = rtl_paragraph(doc)
        p.add_run("☐ " + label)
        if url:
            p.add_run(" → ")
            add_hyperlink(p, "פתח", url)

    add_heading_rtl(doc, "1. פתיחה (5 דק)", 1)
    rows = [
        ("מה הושלם", "Chapters על כל האתר; content-diff; team_50 + team_190 PASS"),
        ("מיזוג", "chapters-home → main @ 6d898a7 (23.6)"),
        ("מה לא קרה", "eyalamit.co.il עדיין legacy — cutover = M7"),
        ("מסגרת", "היום: סקירה + אישורים; אחרי: M5 → M6 → M7"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (topic, say) in enumerate(rows):
        table.rows[i].cells[0].text = topic
        table.rows[i].cells[1].text = say

    add_heading_rtl(doc, "2. סקירה חיה על staging (40 דק)", 1)
    p = rtl_paragraph(doc)
    p.add_run("סדר מומלץ — לחץ על כל שורה:")
    tour = [
        ("בית", "/"),
        ("השיטה", "/method/"),
        ("טיפול", "/treatment/"),
        ("סאונד הילינג", "/sound-healing/"),
        ("שיעורים", "/lessons/"),
        ("אודות", "/eyal-amit/"),
        ("מוקש דהימן", "/eyal-amit/mokesh-dahiman/"),
        ("מוזה + ספרים", "/books/"),
        ("חנות (5 מוצרים)", "/shop/"),
        ("FAQ", "/faq/"),
        ("בלוג", "/blog/"),
        ("צור קשר", "/contact/"),
    ]
    for i, (label, path) in enumerate(tour, 1):
        p = rtl_paragraph(doc)
        p.add_run(f"{i}. {label} — ")
        add_hyperlink(p, staging(path), staging(path))
    p = rtl_paragraph(doc)
    p.add_run("במהלך המעבר: הדגש verbatim, נגישות, wa.me, טפסים — לא תיקונים טכניים.")

    add_heading_rtl(doc, "3. אישורים ממך — לפי עדיפות (30 דק)", 1)

    sections = [
        (
            "עדיפות #1 — 15 הצעות תוכן SEO/GEO",
            [
                ("דף Hub", hub("content-proposals.html")),
                ("מנוף עסקי", "CP-01 פילר נחירות/דום נשימה"),
                ("יעד", "אישור / דחייה / עריכה; ייצוא JSON"),
            ],
        ),
        (
            "#2 — 48 עדויות",
            [
                ("נקודת כניסה", hub("what-we-need.html")),
                ("השלמות I1", hub("materials-intake.html")),
                ("יעד", "אישור אופן הצגה + אצירה"),
            ],
        ),
        (
            "#3 — מותג מושמט (WP-06)",
            [
                ("החלטה", "סטודיו נשימה מעגלית — מושמט או ניסוח חלופי"),
                ("קשור", hub("content-proposals.html") + " (BN-01)"),
            ],
        ),
        (
            "#4 — מדיה",
            [
                ("השלמות", hub("materials-intake.html")),
                ("מדיה ותמונות", hub("media-intake.html")),
                ("חסר", "hero ל-4 שירותים, כריכות, מוקש/גלריות"),
            ],
        ),
        (
            "#5 — משפטי + EN",
            [
                ("EN לאישור", staging("/en/")),
                ("נדרש", "פרטיות / נגישות / תקנון — נוסח או אישור טיוטה"),
            ],
        ),
        (
            "#6 — Clarity (EYL-1)",
            [
                ("מדריך Clarity", hub("files/to-eyal/2026-06-19--clarity-setup-guide/clarity-setup-guide-he.html")),
                ("נדרש", "project_id לפני go-live (GA4 כבר פעיל)"),
            ],
        ),
        (
            "#7 — וואטסאפ",
            [
                ("אישור", "+972 52-482-2842 — materials-intake (WA)"),
            ],
        ),
    ]
    for title, items in sections:
        add_heading_rtl(doc, title, 2)
        for label, val in items:
            p = rtl_paragraph(doc)
            if val.startswith("http"):
                p.add_run(f"• {label}: ")
                add_hyperlink(p, val.replace(HUB, "Hub: ").replace(STAGING, "Staging: ")[:60], val)
            else:
                p.add_run(f"• {label}: {val}")

    add_heading_rtl(doc, "4. סגירה (10 דק)", 1)
    close_rows = [
        ("M5 timeline", "SEO pillar, מדיה, עדויות — מה חוסם?"),
        ("M6", "design-QA responsive — אחרי מדיה"),
        ("M7 יעד", "תאריך cutover ל-eyalamit.co.il"),
        ("ממך השבוע", "proposals + 3 hero + Clarity"),
    ]
    t2 = doc.add_table(rows=len(close_rows), cols=2)
    t2.style = "Table Grid"
    for i, (q, note) in enumerate(close_rows):
        t2.rows[i].cells[0].text = q
        t2.rows[i].cells[1].text = note

    add_heading_rtl(doc, "נקודות מפתח (אם נשאל)", 1)
    for line in [
        "למה לא prod? — M7 מכוון; חוסמים = תוכן/מדיה/אישורים",
        "team_190? — PASS 23.6; F110-01 סגור",
        "תוכן שונה? — לא; verbatim; חריג מותג מאושר",
    ]:
        add_bullet_rtl(doc, line)

    add_heading_rtl(doc, "קישורי Hub — לחיצה", 1)
    hub_links = [
        ("מה נדרש ממך — עדיפית", "what-we-need.html"),
        ("כניסה", "index.html"),
        ("תדריך פגישה", "meeting.html"),
        ("15 הצעות תוכן", "content-proposals.html"),
        ("פערי תוכן", "what-we-need.html"),
        ("השלמות מאייל", "materials-intake.html"),
        ("מדיה ותמונות", "media-intake.html"),
        ("משימות והחלטות", "tasks.html"),
        ("מפת דרכים M1–M7", "roadmap.html"),
        ("קליטת תוכן", "content-intake.html"),
    ]
    for label, path in hub_links:
        p = rtl_paragraph(doc)
        p.add_run(f"{label}: ")
        add_hyperlink(p, hub(path), hub(path))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"OK: {out_path}")


if __name__ == "__main__":
    root = Path(__file__).resolve().parent.parent
    out = root / "_COMMUNICATION/team_00/MEETING-EYAL-CHAPTERS-2026-06-23-CHECKLIST.docx"
    build(out)
