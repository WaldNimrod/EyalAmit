"""Shared RTL + font helpers for Hebrew CEO deliverables (python-docx)."""
import re
from pathlib import Path

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_run_font(run, name: str = "Arial", size: int = 11, bold: bool = False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)
    rPr.insert(0, rFonts)


def add_rtl_paragraph(doc, text: str, bold: bool = False, size: int = 11, heading_level=None):
    if heading_level is not None:
        p = doc.add_heading(text, level=heading_level)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in p.runs:
            set_run_font(run, size=14 if heading_level == 1 else 12, bold=True)
        return p
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_hyperlink(paragraph, text: str, url: str, name: str = "Arial", size: int = 11):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    r_pr.append(r_style)

    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:cs"), name)
    r_pr.append(r_fonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(size * 2))
    r_pr.append(sz)

    sz_cs = OxmlElement("w:szCs")
    sz_cs.set(qn("w:val"), str(size * 2))
    r_pr.append(sz_cs)

    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "0")
    r_pr.append(rtl)

    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def strip_md_links(text: str) -> str:
    return re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", text)


def strip_md_bold(text: str) -> str:
    return re.sub(r"\*\*([^*]+)\*\*", r"\1", text)


def strip_md_code(text: str) -> str:
    return re.sub(r"`([^`]+)`", r"\1", text)


def clean_inline_md(text: str) -> str:
    return strip_md_code(strip_md_links(strip_md_bold(text)))


def add_md_file_as_docx(doc: Document, md_path: Path, *, skip_first_h1: bool = False):
    """Markdown → docx: headings, lists, paragraphs; RTL."""
    raw = md_path.read_text(encoding="utf-8")
    lines = raw.splitlines()
    in_code = False
    skipped_h1 = False
    for line in lines:
        s = line.rstrip()
        if s.strip().startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            if s.strip():
                add_rtl_paragraph(doc, clean_inline_md(s), size=10)
            continue
        if not s.strip():
            continue
        if s.strip() == "---":
            continue
        if s.startswith("# "):
            if skip_first_h1 and not skipped_h1:
                skipped_h1 = True
                continue
            add_rtl_paragraph(doc, clean_inline_md(s[2:].strip()), heading_level=1)
        elif s.startswith("## "):
            add_rtl_paragraph(doc, clean_inline_md(s[3:].strip()), heading_level=2)
        elif s.startswith("### "):
            add_rtl_paragraph(doc, clean_inline_md(s[4:].strip()), heading_level=3)
        elif s.lstrip().startswith("- ") or s.lstrip().startswith("* "):
            indent = len(s) - len(s.lstrip())
            prefix = "  " * min(indent // 2, 4) + "• "
            add_rtl_paragraph(doc, prefix + clean_inline_md(s.lstrip()[2:].strip()))
        elif re.match(r"^\s*\d+\.\s+", s):
            add_rtl_paragraph(doc, clean_inline_md(s.strip()))
        else:
            add_rtl_paragraph(doc, clean_inline_md(s.strip()))
